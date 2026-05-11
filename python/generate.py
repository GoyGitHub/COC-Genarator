#!/usr/bin/env python3
"""
Simple Certificate Generator - Populates DOCX template and saves to Certificates folder
"""
import re
import sys
import sqlite3
import shutil
from pathlib import Path
from datetime import datetime
from docx import Document

# Configuration
BASE_DIR = Path(__file__).parent.parent
DB_PATH = BASE_DIR / 'storage' / 'hrmo_coc.sqlite'  # SAME DATABASE AS WEB FORM
TEMPLATE_PATH = BASE_DIR / 'template' / 'CERTIFICATION-COLLEGE.docx'
CERTIFICATES_DIR = BASE_DIR / 'Certificates'
OUTPUT_DIR = CERTIFICATES_DIR

# Ensure Certificates folder exists
CERTIFICATES_DIR.mkdir(exist_ok=True)


def get_ordinal_suffix(day: int) -> str:
    """Get ordinal suffix for a day (st, nd, rd, th)"""
    if day % 100 in [11, 12, 13]:
        return 'th'
    if day % 10 == 1:
        return 'st'
    if day % 10 == 2:
        return 'nd'
    if day % 10 == 3:
        return 'rd'
    return 'th'


def format_date_long_parts(date_str: str):
    """Return structured date pieces so suffix can be superscripted."""
    try:
        dt = datetime.strptime(date_str, '%Y-%m-%d')
        day = dt.day
        suffix = get_ordinal_suffix(day)
        month = dt.strftime('%B')
        year = dt.year
        return [
            (str(day), False),
            (suffix, True),
            (f" day of {month} {year}", False),
        ]
    except:
        return [(date_str, False)]


def format_issued_date_parts():
    """Return structured issued date without leading 'this' so template controls the prefix."""
    now = datetime.now()
    day = now.day
    suffix = get_ordinal_suffix(day)
    month = now.strftime('%B')
    year = now.year
    return [
        (str(day), False),
        (suffix, True),
        (f" day of {month} {year}", False),
    ]


def get_pronoun(gender: str, female_word: str, male_word: str) -> str:
    """Return pronoun based on gender"""
    return female_word if gender.lower() == 'female' else male_word


def get_intern_data(intern_id: int) -> dict:
    """Get intern data from database"""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM interns WHERE id = ?', (intern_id,))
    row = cursor.fetchone()
    conn.close()
    
    if not row:
        raise Exception(f'Intern ID {intern_id} not found')
    
    return dict(row)


def split_text_segments(text: str, replacements: dict) -> list:
    """Split text into plain and placeholder segments."""
    segments = []
    position = 0
    while position < len(text):
        next_placeholder = None
        next_index = len(text)
        for placeholder in replacements.keys():
            idx = text.find(placeholder, position)
            if idx != -1 and idx < next_index:
                next_index = idx
                next_placeholder = placeholder
        if next_placeholder is None:
            segments.append((text[position:], None))
            break
        if next_index > position:
            segments.append((text[position:next_index], None))
        segments.append((next_placeholder, next_placeholder))
        position = next_index + len(next_placeholder)
    return segments


def copy_run_format(source_run, target_run):
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    if source_run.font.name:
        target_run.font.name = source_run.font.name
    if source_run.font.size:
        target_run.font.size = source_run.font.size
    if source_run.font.color and source_run.font.color.rgb:
        target_run.font.color.rgb = source_run.font.color.rgb
    return target_run


def replace_text_in_paragraph(paragraph, replacements):
    """Replace text in paragraph while preserving run formatting and superscript."""
    for run in list(paragraph.runs):
        if not any(placeholder in run.text for placeholder in replacements.keys()):
            continue

        segments = split_text_segments(run.text, replacements)
        if len(segments) == 1 and segments[0][1] is None:
            continue

        run_text = segments[0][0]
        run_placeholder = segments[0][1]
        last_run = run
        if run_placeholder is None:
            run.text = run_text
            run.font.superscript = False
        else:
            replacement = replacements[run_placeholder]
            if isinstance(replacement, list):
                first_text, first_sup = replacement[0]
                run.text = first_text
                run.font.superscript = first_sup
                for extra_text, extra_sup in replacement[1:]:
                    new_run = paragraph.add_run(extra_text)
                    copy_run_format(run, new_run)
                    new_run.font.superscript = extra_sup
                    last_run._r.addnext(new_run._r)
                    last_run = new_run
            else:
                run.text = replacement
                run.font.superscript = False

        for text_segment, placeholder in segments[1:]:
            if placeholder is None:
                if text_segment == '':
                    continue
                new_run = paragraph.add_run(text_segment)
                copy_run_format(run, new_run)
                new_run.font.superscript = False
                last_run._r.addnext(new_run._r)
                last_run = new_run
                continue

            replacement = replacements[placeholder]
            if isinstance(replacement, list):
                for piece_text, piece_sup in replacement:
                    new_run = paragraph.add_run(piece_text)
                    copy_run_format(run, new_run)
                    new_run.font.superscript = piece_sup
                    last_run._r.addnext(new_run._r)
                    last_run = new_run
            else:
                new_run = paragraph.add_run(replacement)
                copy_run_format(run, new_run)
                new_run.font.superscript = False
                last_run._r.addnext(new_run._r)
                last_run = new_run


def generate_certificate(intern_id: int) -> str:
    """Generate certificate DOCX file"""
    
    # Get intern data
    intern = get_intern_data(intern_id)
    
    # Build course text
    course_text = ''
    if intern.get('course'):
        course_text = intern['course'] + ', '
    
    # Build replacements dict
    replacements = {
        '{{FULL_NAME}}': intern['full_name'].upper(),
        '{{COURSE}}': intern['course'] or '',
        '{{SCHOOL}}': intern['school'].upper(),
        '{{HOURS_RENDERED}}': str(intern['hours_rendered']),
        '{{DEPARTMENT}}': intern['department'],
        '{{START_DATE}}': format_date_long_parts(intern['start_date']),
        '{{END_DATE}}': format_date_long_parts(intern['end_date']),
        '{{ISSUED_DATE}}': format_issued_date_parts(),
        '{{PRONOUN_HER_HIS}}': get_pronoun(intern['gender'], 'her', 'his'),
        '{{CERTIFICATE_ID}}': intern['certificate_id'] or '',
    }
    
    # Copy template to preserve original
    name_cleaned = intern['full_name'].replace(' ', '-')
    output_filename = f"COC-{name_cleaned}.docx"
    output_path = OUTPUT_DIR / output_filename
    
    # Copy template
    shutil.copy2(TEMPLATE_PATH, output_path)
    
    # Load and modify
    doc = Document(output_path)
    
    # Replace in all paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

    # Update bottom date to today's date in MM.DD.YYYY format
    today_bottom_date = datetime.now().strftime('%m.%d.%Y')
    for paragraph in doc.paragraphs:
        if re.match(r'^\d{2}\.\d{2}\.\d{4}$', paragraph.text.strip()):
            if paragraph.runs:
                paragraph.runs[0].text = today_bottom_date
            else:
                paragraph.add_run(today_bottom_date)

    # Save
    doc.save(output_path)
    
    return str(output_filename)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python generate.py <intern_id>')
        sys.exit(1)
    
    try:
        intern_id = int(sys.argv[1])
        filename = generate_certificate(intern_id)
        # Use simple ASCII output to avoid encoding issues
        print('Certificate generated: {}'.format(filename))
        sys.exit(0)
    except Exception as e:
        print('Error: {}'.format(str(e)), file=sys.stderr)
        sys.exit(1)
